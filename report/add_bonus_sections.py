from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
REPORT = BASE / "cloud_course_design_report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def paragraph(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def code_block(doc, text):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, 120, 120, 160, 160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def main():
    doc = Document(REPORT)
    if any("八、附加题实践" in p.text for p in doc.paragraphs):
        print("bonus section already exists")
        return

    doc.add_page_break()
    doc.add_heading("八、附加题实践", level=1)
    paragraph(
        doc,
        "本项目进一步围绕监控系统、CI/CD 流水线和边缘计算专题完成附加题设计。"
        "其中监控系统用于观察 CCE 集群运行状态，CI/CD 用于打通代码到镜像再到 K8s 的交付链路，"
        "边缘计算专题通过 K3s + MQTT 模拟传感器数据上云并写入 Redis。",
    )

    doc.add_heading("8.1 附加题1：Prometheus + Grafana 监控系统", level=2)
    paragraph(
        doc,
        "监控系统采用 kube-prometheus-stack 部署 Prometheus、Grafana、node-exporter 和 kube-state-metrics。"
        "Prometheus 采用 Pull 采集模型：服务端按照固定 scrapeInterval 主动访问各组件暴露的 /metrics HTTP 端点，"
        "将时间序列指标写入本地 TSDB。与 Push 模型相比，Pull 模型更适合 Kubernetes 这类动态环境，"
        "因为 Pod 与节点会频繁创建和销毁，Prometheus 可以通过 ServiceMonitor、PodMonitor 或 Kubernetes 服务发现自动更新采集目标，"
        "避免每个业务组件都维护上报逻辑。Grafana 则负责把 Prometheus 查询语言 PromQL 的结果可视化成折线图、柱状图和仪表盘。",
    )
    add_table(
        doc,
        ["指标", "含义", "在本实验中的用途"],
        [
            [
                "节点 CPU 利用率",
                "节点 CPU 在一段时间内被系统组件与业务 Pod 消耗的比例。",
                "用于判断新增 Spark、监控和 Web 应用后节点是否接近瓶颈。",
            ],
            [
                "Pod 内存使用量",
                "Pod 当前实际使用的工作集内存。",
                "用于观察 backend、redis、spark driver/executor 等组件是否超出预期。",
            ],
            [
                "Pod 重启次数",
                "容器被 kubelet 重启的累计次数。",
                "用于识别 CrashLoopBackOff、OOMKilled 或配置错误导致的不稳定。",
            ],
        ],
    )
    code_block(
        doc,
        "helm upgrade --install cloud-monitor prometheus-community/kube-prometheus-stack "
        "-n monitoring --create-namespace -f monitoring/kube-prometheus-stack-values.yaml\n"
        "kubectl -n monitoring get pods -o wide\n"
        "kubectl -n monitoring get svc",
    )
    paragraph(doc, "验收截图待补：Grafana 节点 CPU 折线图、Pod 内存柱状图、monitoring 命名空间 Pod Running。", bold=True, color="C00000")

    doc.add_heading("8.2 附加题2：CI/CD 流水线", level=2)
    paragraph(
        doc,
        "CI/CD 流水线使用 GitHub Actions 实现。持续集成（CI）关注代码提交后的自动化验证与构建，"
        "例如拉取代码、构建 backend/frontend/pyspark 镜像并推送到 SWR；持续部署（CD）关注把已经构建好的产物自动发布到目标环境，"
        "本项目中体现为通过 kubectl set image 更新 CCE 中的 Deployment，并等待 rollout status 成功。"
        "流水线使用 commit hash 作为镜像 Tag，使每一次部署都能追溯到具体代码版本。"
        "如果线上出现问题，可以根据 Deployment 中的镜像 Tag 快速定位对应提交，必要时回滚到上一版本。",
    )
    paragraph(
        doc,
        "GitOps 的核心思想是以 Git 仓库作为系统期望状态的唯一事实来源。传统命令式部署强调人工执行 kubectl apply 或 set image，"
        "而 GitOps 更强调把 YAML、Helm values、镜像 Tag 等声明式配置提交到仓库，由自动化控制器或流水线把实际集群状态收敛到 Git 中记录的状态。"
        "这样做的好处包括：变更可审计、环境可复现、回滚简单、多人协作冲突更容易通过代码评审发现。"
        "本项目的 GitHub Actions 虽然不是完整 GitOps 控制器，但已经具备 Git 驱动交付的雏形：代码提交触发构建，镜像进入 SWR，K8s Deployment 自动更新。",
    )
    code_block(
        doc,
        ".github/workflows/huawei-swr-cce.yml\n"
        "docker build --provenance=false ...\n"
        "docker push swr.cn-east-3.myhuaweicloud.com/cloud-swjtu/backend:<commit>\n"
        "kubectl -n cloud-course set image deployment/backend backend=<new-image>",
    )
    paragraph(doc, "验收截图待补：GitHub Actions 全部 Passed、Push images 日志、CCE Deployment 镜像 Tag 自动更新。", bold=True, color="C00000")

    doc.add_heading("8.3 附加题3：前沿专题 C-2 K3s + MQTT 边缘计算模拟", level=2)
    paragraph(
        doc,
        "本专题选择“边缘计算模拟：K3s + MQTT”。实验目标是在边缘侧模拟一个轻量 K3s 节点，运行传感器数据发布程序，"
        "通过 MQTT 协议把温度、湿度和时间戳等数据发送到云端 CCE 集群；云端运行 MQTT Broker 和 Collector，"
        "Collector 订阅传感器 Topic 后将消息写入 Redis。该过程模拟了实际物联网场景中“边缘设备采集、消息协议上云、云端存储与后续分析”的链路。",
    )
    paragraph(
        doc,
        "选择 K3s 作为边缘侧运行环境，是因为它是面向资源受限场景的轻量 Kubernetes 发行版，单节点即可运行，"
        "适合部署在工控机、边缘网关或小型虚拟机上。相比完整 Kubernetes，K3s 删除或替换了一些重量级组件，"
        "安装包更小，启动更快，对 CPU 和内存的要求更低。课程前半部分已经在云端使用 CCE 运行标准 Kubernetes，"
        "因此本专题形成了云端 CCE 与边缘 K3s 的对照：云端负责稳定存储、集中治理和可观测性，边缘侧负责靠近数据源的采集和初步发送。",
    )
    paragraph(
        doc,
        "MQTT 协议采用发布/订阅模型。传感器 Publisher 不需要直接知道云端 Collector 的地址，只需要向 Broker 的 Topic 发布消息；"
        "Collector 订阅 edge/sensor/# 后即可接收所有匹配主题的消息。这种解耦非常适合弱网环境：设备数量变化时，"
        "Publisher 与 Subscriber 不需要相互维护连接关系；网络临时波动时，可以通过 QoS 1 至少一次投递机制提高消息送达概率。"
        "本实验中的 sensor_publisher.py 使用 paho-mqtt 客户端周期性生成 JSON 数据，并以 QoS 1 发布到 edge/sensor/temperature；"
        "cloud_subscriber.py 订阅该主题后，将消息包装为带 received_at 的记录写入 Redis List edge:mqtt:events。",
    )
    paragraph(
        doc,
        "该架构的关键挑战是云边协同延迟与可靠性。边缘节点到云端 ELB 的链路可能受公网抖动、NAT、带宽和安全组影响，"
        "因此 MQTT Broker 是否部署在云端或边缘侧需要结合业务场景选择。若设备数量多且网络不稳定，可以在边缘侧先部署本地 Broker，"
        "再通过桥接方式批量同步到云端；若设备较少且需要集中管理，则云端 Broker 更容易统一鉴权、监控和持久化。"
        "本实验采用云端 Broker，是为了复用 CCE、ELB 和 Redis，突出云平台统一承载能力。若继续扩展，可以加入 MQTT 用户名密码、TLS、"
        "消息离线缓存、Redis Stream 或 Kafka，从而提升安全性和可追溯性。",
    )
    paragraph(
        doc,
        "从课程知识角度看，该实验把云计算、容器编排、消息中间件和边缘计算连接起来：K3s/Kubernetes 负责容器运行和声明式管理，"
        "MQTT 负责轻量消息传输，Redis 负责云端状态存储，ELB 负责公网入口。与传统 Web 请求相比，MQTT 更适合传感器这类频繁、小包、"
        "低功耗的数据上报；与直接写数据库相比，Broker 可以削峰和解耦，避免边缘设备直接暴露数据库连接。"
        "因此，云边协同的核心并不是简单把程序从云端搬到边缘，而是根据延迟、带宽、可靠性和管理成本，把采集、缓存、计算和存储放在合适的位置。",
    )
    code_block(
        doc,
        "edge_mqtt/sensor_publisher.py\n"
        "edge_mqtt/cloud_subscriber.py\n"
        "edge_mqtt/k8s-cloud-mqtt.yaml\n"
        "edge_mqtt/k3s-edge-publisher-job.yaml",
    )
    paragraph(doc, "验收截图待补：边缘 Publisher 日志、云端 Collector 日志、Redis 中 edge:mqtt:events 数据。", bold=True, color="C00000")

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
