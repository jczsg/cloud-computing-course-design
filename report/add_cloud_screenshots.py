from pathlib import Path

from PIL import Image, ImageChops
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
REPORT = BASE / "cloud_course_design_report.docx"
SCREENSHOT_DIR = BASE / "screenshots"
CROPPED_DIR = BASE / "outputs" / "cloud_acceptance_screenshots"


CAPTIONS = [
    "图 6-1 CCE 集群节点验证：4 个 Worker 节点均为 Ready，Kubernetes 版本为 v1.35.3。",
    "图 6-2 cloud-course 命名空间运行状态：backend 与 redis Pod 均为 Running。",
    "图 6-3 Service 暴露结果：backend-svc 绑定公网 ELB 地址 1.94.27.202。",
    "图 6-4 后端公网接口验收：访问 /api/ping 返回 status=ok。",
    "图 6-5 Redis 持久化卷验收：redis-data-pvc 已 Bound，容量 10Gi。",
    "图 6-6 Redis 持久化验证：写入 testkey 后删除 Pod，新 Pod 仍能读取 hello。",
    "图 6-7 ConfigMap Volume 验证：Nginx default.conf 已从 ConfigMap 挂载到容器内。",
    "图 6-8 HPA 排查结果：backend-hpa 已创建，但集群 Metrics API 不可用，无法读取 CPU 指标。",
    "图 6-9 Deployment 手动扩容验证：backend 从 1 副本扩到 2 副本并成功调度到不同节点。",
    "图 6-10 Deployment 手动缩容验证：backend 缩回 1 副本后仍保持 Running。",
    "图 6-11 Spark Operator 镜像修复：controller 镜像切换至 SWR 后 Pod Running。",
    "图 6-12 Spark 直接 Pod 作业验收：spark-wordcount-direct 执行完成，Pod 状态为 Completed。",
    "图 6-13 Spark WordCount 日志：输出 Top 10 words，说明 PySpark 作业执行成功。",
    "图 6-14 华为云 SWR 控制台截图：frontend 镜像已上传至 cloud-swjtu 组织。",
    "图 6-15 华为云 SWR 控制台截图：backend 镜像已上传至 cloud-swjtu 组织。",
    "图 6-16 华为云 SWR 控制台截图：pyspark 镜像已上传至 cloud-swjtu 组织。",
    "图 6-17 华为云 SWR 控制台截图：spark-operator-controller 镜像已上传至 cloud-swjtu 组织。",
]


def crop_white_margin(src: Path, dst: Path) -> Path:
    image = Image.open(src).convert("RGB")
    bg = Image.new("RGB", image.size, (255, 255, 255))
    diff = ImageChops.difference(image, bg)
    bbox = diff.getbbox()
    if not bbox:
        image.save(dst)
        return dst

    left, top, right, bottom = bbox
    pad = 16
    left = max(0, left - pad)
    top = max(0, top - pad)
    right = min(image.width, right + pad)
    bottom = min(image.height, bottom + pad)
    image.crop((left, top, right, bottom)).save(dst)
    return dst


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def main():
    if not REPORT.exists():
        raise FileNotFoundError(REPORT)
    screenshots = sorted(SCREENSHOT_DIR.glob("*.png"))
    if len(screenshots) != len(CAPTIONS):
        raise RuntimeError(f"期望 {len(CAPTIONS)} 张截图，实际找到 {len(screenshots)} 张。")

    CROPPED_DIR.mkdir(parents=True, exist_ok=True)
    cropped = []
    for idx, src in enumerate(screenshots, 1):
        dst = CROPPED_DIR / f"{idx:02d}-{src.stem}.png"
        cropped.append(crop_white_margin(src, dst))

    doc = Document(REPORT)

    for p in doc.paragraphs:
        if p.text.startswith("截图占位："):
            p.text = p.text.replace("截图占位：", "验收截图：") + "（实际截图见第六章）"

    doc.add_page_break()
    doc.add_heading("六、云端验收截图汇总", level=1)
    doc.add_paragraph(
        "本章汇总华为云 CCE、SWR、Kubernetes 命令行和 Spark 作业的实际验收截图，"
        "用于对应正文中集群部署、服务暴露、存储持久化、ConfigMap、HPA 和 Spark 运行结果。"
    )

    for path, caption in zip(cropped, CAPTIONS):
        doc.add_picture(str(path), width=Inches(6.25))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(doc, caption)

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
