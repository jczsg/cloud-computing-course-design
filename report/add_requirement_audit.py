from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parents[1]
REPORT = BASE / "cloud_course_design_report.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "bottom": bottom, "start": start, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, "F2F4F7")
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    return table


def paragraph(doc, text, bold=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    return p


def main():
    doc = Document(REPORT)

    # Avoid appending the audit repeatedly when the script is rerun.
    if any("七、任务书逐项对照检查" in p.text for p in doc.paragraphs):
        print("audit section already exists")
        return

    doc.add_page_break()
    doc.add_heading("七、任务书逐项对照检查", level=1)
    paragraph(
        doc,
        "本节按照《课程设计任务书》的评分项对当前成果进行自查。状态分为：已满足、基本满足、部分满足。"
        "其中“部分满足”并不表示无法提交，而是说明当前云环境或截图证据与任务书满分要求存在差距，已在报告中给出原因说明或替代验证。",
    )

    add_table(
        doc,
        ["任务书评分项", "当前完成情况", "状态"],
        [
            [
                "任务1 应用容器化：Dockerfile 多阶段、自选 Python 包、首页含学号姓名、本地联调、SWR 镜像截图。",
                "代码中已提供 backend/frontend Dockerfile、requests 依赖、前端首页和 SWR 镜像截图；本地联调有命令与证据文件，若教师严格要求“docker compose 窗口截图”，可额外补一张。",
                "基本满足",
            ],
            [
                "任务2 CCE 集群搭建：至少 2 个 Worker Ready，VERSION >= 1.27。",
                "云端截图显示 4 个 Worker 节点 Ready，版本 v1.35.3，满足版本与节点要求。",
                "已满足",
            ],
            [
                "任务3 应用部署：后端 Deployment 副本=2，Redis 副本=1，resources、ConfigMap、Secret、LoadBalancer、/api/ping。",
                "YAML 已包含 backend replicas=2、Redis=1、resources、ConfigMap/Secret；云端已验证 backend-svc 公网 /api/ping 返回 status=ok。",
                "已满足",
            ],
            [
                "任务4 持久化存储：storageClassName=csi-disk，PVC Bound，Redis 写入、删 Pod、重建后仍可 GET。",
                "YAML 已修正为 csi-disk；云端截图包含 PVC Bound、SET testkey、删 Pod、重建后 GET hello。",
                "已满足",
            ],
            [
                "任务5 ConfigMap Volume：Nginx 配置以 Volume 挂载，exec 查看文件更新，并说明 Volume 与 envFrom 差异。",
                "报告已说明两种方式差异；云端使用同一 Nginx 镜像挂载同一 nginx-config 的检查 Pod 验证 default.conf，能证明 ConfigMap Volume 内容生效。严格按题面应补“frontend Pod 内 exec”截图。",
                "基本满足",
            ],
            [
                "任务6 HPA：metrics 可用、ab 压测、Pod 从 1 扩到 2+、停压后缩回 1，并分析延迟/冷却/降本。",
                "backend-hpa 已创建，手动扩缩容验证成功；但云端 Metrics API 不可用，HPA 无法读取 pods.metrics.k8s.io。任务书允许未触发时提供 describe hpa 排查日志并分析原因，报告已覆盖。",
                "基本满足",
            ],
            [
                "Spark A-0：Spark Operator 提交 SparkApplication，Driver 与 Executor Pod Completed 截图。",
                "Spark Operator 已修复 Running；SparkApplication 因资源 requests 持续 Pending。报告补充同一 PySpark 镜像的 spark-wordcount-direct Pod Completed 和日志作为替代验证。",
                "部分满足",
            ],
            [
                "Spark A-1：加载数据到 DataFrame，打印 Schema 和前 5 行，缺失值比例，2 种缺失处理，清洗前后行数和基本统计。",
                "代码 spark/douban_spark_analysis.py 已实现 schema、show(5)、缺失比例、dropna/fillna、清洗前后行数和 describe；报告表格使用本地复现结果。缺少 Spark on K8s 实际运行截图。",
                "基本满足",
            ],
            [
                "Spark A-2：至少 4 个统计查询，包含 GROUP BY、Top-N、时间趋势、JOIN 或窗口函数；每个查询截图和不少于 50 字分析。",
                "代码和报告已覆盖 5 个查询，文字分析充足，报告中有结果表和图；但缺少每个 Spark 查询在 K8s/CloudShell 中输出的原始截图。",
                "基本满足",
            ],
            [
                "Spark A-3：Pandas vs PySpark executor=1/2 执行时间，对比图，Amdahl 量化分析。",
                "当前只有 Pandas 实测时间和 Amdahl 原理分析；PySpark 1/2 executor 时间与对比图尚未实测。此项是当前最明显缺口。",
                "部分满足",
            ],
            [
                "报告质量：封面、环境信息、任务记录、截图、性能图表、总结不少于 200 字、附录代码或仓库链接。",
                "报告结构完整、总结超过 200 字、截图清晰；封面建议补班级，最终提交需导出 PDF，并附代码仓库链接或项目压缩包。",
                "基本满足",
            ],
        ],
    )

    paragraph(doc, "当前最建议补强的 3 项：", bold=True, color="C00000")
    for item in [
        "补 Spark 数据分析实际运行截图：schema、前 5 行、缺失比例、4 个查询输出。",
        "补 Spark 性能对比：executorInstances=1 和 2 的执行时间，并生成对比图。",
        "导出 PDF，并在封面补班级；如教师非常严格，再补 frontend Pod 内 exec default.conf 的截图。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

    doc.save(REPORT)
    print(REPORT)


if __name__ == "__main__":
    main()
